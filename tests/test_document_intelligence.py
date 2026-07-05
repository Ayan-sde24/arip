"""Unit tests for the Document Intelligence Pipeline (TASK-004)."""

import io
from dataclasses import FrozenInstanceError
from datetime import UTC, datetime
from uuid import uuid4

import fitz  # type: ignore[import-untyped]
import pytest
from docx import Document as DocxDocument  # type: ignore[import-untyped]

from app.domain.entities.document import Document, DocumentStatus, DocumentType
from app.domain.entities.document_content import DocumentContent
from app.domain.entities.page import Page
from app.domain.entities.text_block import TextBlock
from app.infrastructure.parser.document_reader import DocumentIntelligencePipeline
from app.infrastructure.parser.docx_reader import DocxReader
from app.infrastructure.parser.exceptions import (
    CorruptedDocxError,
    EmptyDocumentError,
    ReaderFailureError,
    UnreadablePDFError,
)
from app.infrastructure.parser.layout_analyzer import LayoutAnalyzer
from app.infrastructure.parser.models import RawDocumentData, RawPageData
from app.infrastructure.parser.pdf_reader import PdfReader
from app.infrastructure.parser.utils import clean_text, compute_statistics

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_document(extension: str = "pdf") -> Document:
    """Return a minimal :class:`Document` domain entity for testing."""
    now = datetime.now(UTC)
    return Document(
        document_id=uuid4(),
        document_type=DocumentType.RESUME,
        original_filename=f"test.{extension}",
        stored_filename=f"{uuid4()}.{extension}",
        mime_type=(
            "application/pdf"
            if extension == "pdf"
            else (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            )
        ),
        extension=extension,
        checksum="abc123",
        size=1024,
        created_at=now,
        updated_at=now,
        status=DocumentStatus.UPLOADED,
    )


def _make_pdf_bytes(text: str = "Hello World from ARIP.") -> bytes:
    """Return valid minimal PDF bytes containing ``text`` using PyMuPDF."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    pdf_bytes: bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    """Return valid DOCX bytes with the given paragraph strings."""
    doc = DocxDocument()
    for para in paragraphs:
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Domain entity tests
# ---------------------------------------------------------------------------


def test_text_block_creation_and_immutability() -> None:
    """TextBlock is created correctly and is immutable."""
    tb = TextBlock(
        text="Engineer with 5 years experience",
        page=1,
        block_index=0,
        reading_order=0,
        bbox=(10.0, 20.0, 200.0, 35.0),
    )
    assert tb.text == "Engineer with 5 years experience"
    assert tb.page == 1
    assert tb.block_index == 0
    assert tb.reading_order == 0
    assert tb.bbox == (10.0, 20.0, 200.0, 35.0)

    with pytest.raises(FrozenInstanceError):
        tb.text = "Changed"  # type: ignore[misc]


def test_text_block_optional_bbox() -> None:
    """TextBlock bbox defaults to None when not supplied."""
    tb = TextBlock(text="No bbox", page=2, block_index=1, reading_order=5)
    assert tb.bbox is None


def test_page_creation_and_immutability() -> None:
    """Page is created with text_blocks and defaults to empty tables/images."""
    tb = TextBlock(text="Block", page=1, block_index=0, reading_order=0)
    pg = Page(page_number=1, text_blocks=[tb], raw_text="Block")
    assert pg.page_number == 1
    assert len(pg.text_blocks) == 1
    assert pg.raw_text == "Block"
    assert pg.tables == []
    assert pg.images == []

    with pytest.raises(FrozenInstanceError):
        pg.page_number = 2  # type: ignore[misc]


def test_document_content_creation_and_immutability() -> None:
    """DocumentContent stores all fields correctly and is immutable."""
    doc = _make_document()
    tb = TextBlock(text="Sample", page=1, block_index=0, reading_order=0)
    pg = Page(page_number=1, text_blocks=[tb], raw_text="Sample")

    dc = DocumentContent(
        document=doc,
        pages=[pg],
        raw_text="Sample",
        clean_text="Sample",
        metadata={"author": "ARIP"},
        language="en",
        statistics={"page_count": 1, "word_count": 1},
    )
    assert dc.document == doc
    assert len(dc.pages) == 1
    assert dc.raw_text == "Sample"
    assert dc.clean_text == "Sample"
    assert dc.metadata["author"] == "ARIP"
    assert dc.language == "en"

    with pytest.raises(FrozenInstanceError):
        dc.language = "fr"  # type: ignore[misc]


# ---------------------------------------------------------------------------
# Utils tests
# ---------------------------------------------------------------------------


def test_clean_text_strips_control_chars() -> None:
    """clean_text removes control characters and normalises whitespace."""
    raw = "Hello\x00World\x01\n\nTest   spacing"
    result = clean_text(raw)
    assert "\x00" not in result
    assert "\x01" not in result
    assert "HelloWorld" in result or "Hello" in result


def test_clean_text_collapses_multiple_newlines() -> None:
    """clean_text collapses 3+ consecutive newlines to 2."""
    raw = "Line1\n\n\n\n\nLine2"
    result = clean_text(raw)
    assert "\n\n\n" not in result
    assert "Line1" in result
    assert "Line2" in result


def test_compute_statistics_word_count() -> None:
    """compute_statistics returns correct word and char counts."""
    text = "Python FastAPI Clean Architecture"
    stats = compute_statistics(raw_text=text, page_count=1)
    assert stats.word_count == 4
    assert stats.char_count == len(text)
    assert stats.page_count == 1
    assert stats.avg_words_per_page == 4.0


def test_compute_statistics_zero_pages() -> None:
    """compute_statistics handles zero page count without division error."""
    stats = compute_statistics(raw_text="some text", page_count=0)
    assert stats.avg_words_per_page == 0.0


# ---------------------------------------------------------------------------
# PDF Reader tests
# ---------------------------------------------------------------------------


def test_pdf_reader_valid_document() -> None:
    """PdfReader successfully reads a valid PDF and extracts text."""
    content = _make_pdf_bytes("ARIP Resume Intelligence Platform")
    reader = PdfReader()
    raw = reader.read(content=content)

    assert raw.page_count >= 1
    all_text = "\n".join(p.raw_text for p in raw.pages)
    assert "ARIP" in all_text


def test_pdf_reader_page_count() -> None:
    """PdfReader returns the correct number of pages."""
    content = _make_pdf_bytes("Single page PDF")
    reader = PdfReader()
    raw = reader.read(content=content)
    assert raw.page_count == 1


def test_pdf_reader_blocks_have_text() -> None:
    """PdfReader extracts non-empty text blocks from the page."""
    content = _make_pdf_bytes("Block text content")
    reader = PdfReader()
    raw = reader.read(content=content)
    assert len(raw.pages) > 0
    all_blocks = [b for p in raw.pages for b in p.blocks]
    assert len(all_blocks) > 0
    assert all(b["text"] for b in all_blocks)


def test_pdf_reader_extract_text_method() -> None:
    """PdfReader.extract_text returns a non-empty string."""
    content = _make_pdf_bytes("Extract text test")
    reader = PdfReader()
    text = reader.extract_text(content=content)
    assert isinstance(text, str)
    assert "Extract" in text


def test_pdf_reader_corrupted_bytes_raises_error() -> None:
    """PdfReader raises UnreadablePDFError for non-PDF bytes."""
    reader = PdfReader()
    with pytest.raises(UnreadablePDFError):
        reader.read(content=b"this is not a pdf at all!!!")


def test_pdf_reader_empty_pdf_raises_error() -> None:
    """PdfReader raises EmptyDocumentError for a PDF with no text."""
    doc = fitz.open()
    doc.new_page()  # blank page — no text
    empty_pdf = doc.tobytes()
    doc.close()

    reader = PdfReader()
    with pytest.raises(EmptyDocumentError):
        reader.read(content=empty_pdf)


def test_pdf_reader_extract_metadata() -> None:
    """PdfReader.extract_metadata returns a dict (possibly empty for test PDFs)."""
    content = _make_pdf_bytes("Metadata test")
    reader = PdfReader()
    meta = reader.extract_metadata(content=content)
    assert isinstance(meta, dict)


# ---------------------------------------------------------------------------
# DOCX Reader tests
# ---------------------------------------------------------------------------


def test_docx_reader_valid_document() -> None:
    """DocxReader successfully reads a valid DOCX and extracts text."""
    content = _make_docx_bytes(["Alice Smith", "Senior Python Engineer", "5 years exp"])
    reader = DocxReader()
    raw = reader.read(content=content)

    assert raw.page_count >= 1
    all_text = "\n".join(p.raw_text for p in raw.pages)
    assert "Alice" in all_text
    assert "Python" in all_text


def test_docx_reader_page_count_single_page() -> None:
    """DocxReader returns at least 1 page when no page breaks are present."""
    content = _make_docx_bytes(["paragraph one", "paragraph two"])
    reader = DocxReader()
    raw = reader.read(content=content)
    assert raw.page_count == 1


def test_docx_reader_blocks_populated() -> None:
    """DocxReader populates blocks list with paragraph entries."""
    content = _make_docx_bytes(["Block A", "Block B", "Block C"])
    reader = DocxReader()
    raw = reader.read(content=content)
    all_blocks = [b for p in raw.pages for b in p.blocks]
    assert len(all_blocks) == 3
    texts = [b["text"] for b in all_blocks]
    assert "Block A" in texts


def test_docx_reader_extract_text_method() -> None:
    """DocxReader.extract_text returns a non-empty string."""
    content = _make_docx_bytes(["Extract this text"])
    reader = DocxReader()
    text = reader.extract_text(content=content)
    assert "Extract" in text


def test_docx_reader_corrupted_bytes_raises_error() -> None:
    """DocxReader raises CorruptedDocxError for non-DOCX bytes."""
    reader = DocxReader()
    with pytest.raises(CorruptedDocxError):
        reader.read(content=b"not a docx file at all")


def test_docx_reader_empty_document_raises_error() -> None:
    """DocxReader raises EmptyDocumentError for a DOCX with no text."""
    doc = DocxDocument()
    buf = io.BytesIO()
    doc.save(buf)
    empty_docx = buf.getvalue()

    reader = DocxReader()
    with pytest.raises(EmptyDocumentError):
        reader.read(content=empty_docx)


def test_docx_reader_extract_metadata() -> None:
    """DocxReader.extract_metadata returns a dict."""
    content = _make_docx_bytes(["Metadata paragraph"])
    reader = DocxReader()
    meta = reader.extract_metadata(content=content)
    assert isinstance(meta, dict)


# ---------------------------------------------------------------------------
# Layout Analyzer tests
# ---------------------------------------------------------------------------


def test_layout_analyzer_assigns_reading_order() -> None:
    """LayoutAnalyzer assigns monotonically increasing reading_order indices."""
    raw = RawDocumentData(
        pages=[
            RawPageData(
                page_number=1,
                raw_text="Block 1\nBlock 2",
                blocks=[
                    {"text": "Block 1", "bbox": (0, 0, 100, 20), "block_index": 0},
                    {"text": "Block 2", "bbox": (0, 25, 100, 45), "block_index": 1},
                ],
            ),
            RawPageData(
                page_number=2,
                raw_text="Block 3",
                blocks=[
                    {"text": "Block 3", "bbox": (0, 0, 100, 20), "block_index": 0},
                ],
            ),
        ]
    )
    analyzer = LayoutAnalyzer()
    pages = analyzer.analyze(raw=raw)

    assert len(pages) == 2
    orders = [tb.reading_order for p in pages for tb in p.text_blocks]
    assert orders == list(range(len(orders))), "reading_order must be monotonic"


def test_layout_analyzer_preserves_page_numbers() -> None:
    """LayoutAnalyzer preserves page numbers from raw data."""
    raw = RawDocumentData(
        pages=[
            RawPageData(
                page_number=3,
                raw_text="Page 3 content",
                blocks=[{"text": "Page 3 content", "bbox": None, "block_index": 0}],
            ),
        ]
    )
    analyzer = LayoutAnalyzer()
    pages = analyzer.analyze(raw=raw)
    assert pages[0].page_number == 3


def test_layout_analyzer_handles_none_bbox() -> None:
    """LayoutAnalyzer sets bbox=None for blocks without coordinate data (DOCX)."""
    raw = RawDocumentData(
        pages=[
            RawPageData(
                page_number=1,
                raw_text="No bbox",
                blocks=[{"text": "No bbox", "bbox": None, "block_index": 0}],
            )
        ]
    )
    analyzer = LayoutAnalyzer()
    pages = analyzer.analyze(raw=raw)
    assert pages[0].text_blocks[0].bbox is None


# ---------------------------------------------------------------------------
# Pipeline integration tests
# ---------------------------------------------------------------------------


def test_pipeline_pdf_end_to_end() -> None:
    """Full pipeline run on a valid PDF produces a DocumentContent."""
    doc = _make_document("pdf")
    pdf_bytes = _make_pdf_bytes("Full pipeline test for ARIP")
    pipeline = DocumentIntelligencePipeline()
    result = pipeline.run(document=doc, content_bytes=pdf_bytes)

    assert isinstance(result, DocumentContent)
    assert result.document == doc
    assert len(result.pages) >= 1
    assert "ARIP" in result.raw_text
    assert result.statistics["page_count"] >= 1
    assert result.statistics["word_count"] > 0


def test_pipeline_docx_end_to_end() -> None:
    """Full pipeline run on a valid DOCX produces a DocumentContent."""
    doc = _make_document("docx")
    docx_bytes = _make_docx_bytes(
        ["Jane Doe", "Machine Learning Engineer", "Skills: Python, PyTorch"]
    )
    pipeline = DocumentIntelligencePipeline()
    result = pipeline.run(document=doc, content_bytes=docx_bytes)

    assert isinstance(result, DocumentContent)
    assert "Jane" in result.raw_text
    assert result.statistics["word_count"] > 0


def test_pipeline_unsupported_extension_raises_error() -> None:
    """Pipeline raises ReaderFailureError for unknown extensions."""
    doc = _make_document("txt")
    pipeline = DocumentIntelligencePipeline()
    with pytest.raises(ReaderFailureError, match="No reader registered"):
        pipeline.run(document=doc, content_bytes=b"plain text")


def test_pipeline_corrupted_pdf_raises_error() -> None:
    """Pipeline propagates UnreadablePDFError for corrupted PDF bytes."""
    doc = _make_document("pdf")
    pipeline = DocumentIntelligencePipeline()
    with pytest.raises(UnreadablePDFError):
        pipeline.run(document=doc, content_bytes=b"not a real pdf")


def test_pipeline_empty_pdf_raises_error() -> None:
    """Pipeline propagates EmptyDocumentError for a blank PDF."""
    fitz_doc = fitz.open()
    fitz_doc.new_page()
    blank_pdf = fitz_doc.tobytes()
    fitz_doc.close()

    doc = _make_document("pdf")
    pipeline = DocumentIntelligencePipeline()
    with pytest.raises(EmptyDocumentError):
        pipeline.run(document=doc, content_bytes=blank_pdf)


def test_pipeline_document_content_is_immutable() -> None:
    """DocumentContent produced by the pipeline is frozen/immutable."""
    doc = _make_document("pdf")
    pdf_bytes = _make_pdf_bytes("Immutable test")
    pipeline = DocumentIntelligencePipeline()
    result = pipeline.run(document=doc, content_bytes=pdf_bytes)

    with pytest.raises(FrozenInstanceError):
        result.language = "fr"  # type: ignore[misc]


def test_pipeline_layout_reading_order_preserved() -> None:
    """Pipeline preserves reading order across multiple pages."""
    doc = _make_document("pdf")
    # Create a two-page PDF
    fitz_doc = fitz.open()
    page1 = fitz_doc.new_page()
    page1.insert_text((72, 72), "First page content here")
    page2 = fitz_doc.new_page()
    page2.insert_text((72, 72), "Second page content here")
    two_page_pdf = fitz_doc.tobytes()
    fitz_doc.close()

    pipeline = DocumentIntelligencePipeline()
    result = pipeline.run(document=doc, content_bytes=two_page_pdf)

    assert result.statistics["page_count"] == 2
    all_blocks = [tb for pg in result.pages for tb in pg.text_blocks]
    orders = [tb.reading_order for tb in all_blocks]
    assert orders == sorted(orders), "reading_order must be monotonically increasing"
