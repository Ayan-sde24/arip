"""Integration tests for the complete Document Intelligence Pipeline (TICKET-005.6)."""

import io
from datetime import UTC, datetime
from uuid import uuid4

import fitz  # type: ignore[import-untyped]
from docx import Document as DocxDocument  # type: ignore[import-untyped]

from app.application.document_analysis.document_pipeline import DocumentPipeline
from app.domain.entities.document import Document, DocumentStatus, DocumentType
from app.domain.entities.section_type import SectionType


def _make_document(extension: str = "pdf") -> Document:
    """Return a minimal Document domain entity for testing."""
    now = datetime.now(UTC)
    return Document(
        document_id=uuid4(),
        document_type=DocumentType.RESUME,
        original_filename=f"test.{extension}",
        stored_filename=f"{uuid4()}.{extension}",
        mime_type=(
            "application/pdf"
            if extension == "pdf"
            else (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            )
        ),
        extension=extension,
        checksum="abc123456",
        size=2048,
        created_at=now,
        updated_at=now,
        status=DocumentStatus.UPLOADED,
    )


def _make_pdf_bytes(lines: list[str]) -> bytes:
    """Return valid PDF bytes containing text lines using PyMuPDF."""
    doc = fitz.open()
    page = doc.new_page()
    y = 50
    for line in lines:
        rect = fitz.Rect(50, y, 500, y + 40)
        page.insert_textbox(rect, line)
        y += 100
    pdf_bytes: bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    """Return valid DOCX bytes with the given paragraph strings."""
    doc = DocxDocument()
    for para in paragraphs:
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_pipeline_valid_pdf_resume() -> None:
    """Test pipeline run on a valid PDF resume with multiple standard sections."""
    doc = _make_document("pdf")
    pdf_bytes = _make_pdf_bytes(
        [
            "Alice Smith (alice@example.com)",
            "Software Engineer at ACME Corp.",
            "Experience",
            "Senior Engineer at ACME Corp since 2020.",
            "Education",
            "BS in Computer Science from Stanford University.",
            "Technical Skills",
            "Python, Rust, FastAPI, AWS, Docker.",
        ]
    )

    pipeline = DocumentPipeline()
    result = pipeline.run(document=doc, content_bytes=pdf_bytes)

    assert result.success is True
    assert len(result.errors) == 0
    assert result.cir is not None

    cir = result.cir
    # Expect 4 sections:
    # 1. Contact (leading section: Alice Smith, Software Engineer)
    # 2. Experience
    # 3. Education
    # 4. Skills
    assert len(cir.sections) == 4
    assert cir.sections[0].section_type == SectionType.CONTACT
    assert cir.sections[1].section_type == SectionType.EXPERIENCE
    assert cir.sections[2].section_type == SectionType.EDUCATION
    assert cir.sections[3].section_type == SectionType.SKILLS

    assert cir.statistics.total_pages == 1
    assert cir.statistics.total_sections == 4
    assert cir.statistics.total_text_blocks == 8


def test_pipeline_valid_docx_resume() -> None:
    """Test pipeline run on a valid DOCX resume with multiple standard sections."""
    doc = _make_document("docx")
    docx_bytes = _make_docx_bytes(
        [
            "Bob Jones (bob@example.com)",
            "Summary",
            "Results-oriented project manager.",
            "Experience",
            "Managed multiple cloud migration projects.",
        ]
    )

    pipeline = DocumentPipeline()
    result = pipeline.run(document=doc, content_bytes=docx_bytes)

    assert result.success is True
    assert len(result.errors) == 0
    assert result.cir is not None

    cir = result.cir
    assert len(cir.sections) == 3
    assert cir.sections[0].section_type == SectionType.CONTACT
    assert cir.sections[1].section_type == SectionType.SUMMARY
    assert cir.sections[2].section_type == SectionType.EXPERIENCE

    assert cir.statistics.total_sections == 3


def test_pipeline_empty_resume() -> None:
    """Test pipeline handling of a blank resume (no text)."""
    # Create empty PDF doc (no text blocks inserted)
    fitz_doc = fitz.open()
    fitz_doc.new_page()
    empty_pdf = fitz_doc.tobytes()
    fitz_doc.close()

    doc = _make_document("pdf")
    pipeline = DocumentPipeline()
    result = pipeline.run(document=doc, content_bytes=empty_pdf)

    # Empty document causes reader stage to raise EmptyDocumentError
    assert result.success is False
    assert len(result.errors) == 1
    err_msg = result.errors[0]
    assert "EmptyDocumentError" in err_msg or "no extractable text" in err_msg
    assert result.cir is None


def test_pipeline_resume_missing_headings() -> None:
    """Test pipeline run on a document containing prose but no structural headings."""
    doc = _make_document("pdf")
    pdf_bytes = _make_pdf_bytes(
        [
            "This is a prose document without clear section headings.",
            "All sentences end with periods and are part of standard text.",
            "Thus, it should be processed as a single document-wide section.",
        ]
    )

    pipeline = DocumentPipeline()
    result = pipeline.run(document=doc, content_bytes=pdf_bytes)

    assert result.success is True
    assert result.cir is not None
    assert len(result.cir.sections) == 1

    # Heading-less boundary maps to SectionType.CONTACT (or other leading default)
    assert result.cir.sections[0].section_type == SectionType.CONTACT
    assert result.cir.sections[0].title is None

    # Warning list should contain a notice about heading absence
    assert any("No headings detected" in w for w in result.warnings)


def test_pipeline_resume_duplicate_headings() -> None:
    """Test pipeline run on a document with multiple sections of the same type."""
    doc = _make_document("pdf")
    pdf_bytes = _make_pdf_bytes(
        [
            "Education",
            "MS in Engineering",
            "Education",
            "BS in Science",
        ]
    )

    pipeline = DocumentPipeline()
    result = pipeline.run(document=doc, content_bytes=pdf_bytes)

    assert result.success is True
    assert result.cir is not None
    cir = result.cir

    # 2 duplicate headings (Education) lead to 2 sections
    assert len(cir.sections) == 2
    assert cir.sections[0].section_type == SectionType.EDUCATION
    assert cir.sections[1].section_type == SectionType.EDUCATION
    assert cir.sections[0].id != cir.sections[1].id


def test_pipeline_resume_unknown_sections() -> None:
    """Test pipeline run on a document containing non-keyword custom headings."""
    doc = _make_document("pdf")
    pdf_bytes = _make_pdf_bytes(
        [
            "Alice Smith (alice@example.com)",
            "UNIQUE MISCELLANEOUS HEADERS",
            "This content belongs under a non-keyword heading.",
            "It should map to SectionType.OTHER.",
        ]
    )

    pipeline = DocumentPipeline()
    result = pipeline.run(document=doc, content_bytes=pdf_bytes)

    assert result.success is True
    assert result.cir is not None
    cir = result.cir

    assert len(cir.sections) == 2  # leading unheaded contact block + OTHER
    assert cir.sections[1].section_type == SectionType.OTHER
    assert cir.sections[1].title == "UNIQUE MISCELLANEOUS HEADERS"
